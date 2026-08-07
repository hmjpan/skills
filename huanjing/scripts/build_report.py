#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report.py - 生成最终气候健康风险预警报告

输入: compute_indices.py 输出的指标 dict
输出:
  1. 完整报告 JSON (report.json) - 含全部指标、证据链、预警文案、不确定性
  2. 交互式风险地图 (map.html, folium, 可选依赖)
  3. Markdown 版报告 (report.md, 适合演示/传播)

用法:
  python build_report.py indices.json --city 上海市 --out-dir ./output
"""
import argparse
import json
import math
import os
import sys
from datetime import datetime

# 预警文案模板（依据风险等级与风险源；数值由调用方从指标中填充，禁止编造）
# 模板结构参考: WHO热浪预警指南 + 参考项目 alert_templates.md
# 每篇含 4 段: 风险结论 -> 为什么危险 -> 具体行动 -> 何时解除/更新
WARNING_TEMPLATES_HEAT = {
    "极高风险": (
        "【高温健康红色预警-紧急】{city}在{period}健康风险极高(评分{score}/5)，"
        "请停止一切非必要户外活动。最高气温达{tmax}C，体感温度(热指数)可达{hi}C，"
        "已持续高温{hw_days}天。\n"
        "为什么紧急: 极端高温可危及生命，热射病风险显著升高。"
        "高温高湿环境下人体散热机制失效，核心体温可能快速升至40C以上。\n"
        "怎么做: 1)停止户外作业与剧烈运动，学校/工地建议停课停工；"
        "2)开放全部避暑纳凉点，对独居老人实施专人看护；"
        "3)发现有人意识模糊、皮肤灼热无汗:立即拨打120并物理降温。\n"
        "预警生效至本周期结束，请互相转告并关注更新。"
    ),
    "高风险": (
        "【高温健康橙色预警】{city}在{period}健康风险较高(评分{score}/5)，请减少外出。"
        "最高气温预计{tmax}C，体感温度{hi}C，已连续{hw_days}天高温。\n"
        "为什么危险: 长时间暴露可能中暑，老年人、儿童、户外工作者、"
        "心脑血管及呼吸系统慢性病患者风险最大(Gasparrini 2015 Lancet)。\n"
        "怎么做: 1)尽量待在室内，开空调或去纳凉点；"
        "2)必须外出避开11-16点，戴帽、带水、结伴；"
        "3)家里有老人每天至少探望或电话一次；"
        "4)出现头晕、恶心、不出汗:立刻到阴凉处，严重时拨打120。\n"
        "预计本周期后缓解，请关注更新。"
    ),
    "中风险": (
        "【高温健康黄色提示】{city}在{period}存在中度高温健康风险(评分{score}/5)。"
        "最高气温{tmax}C，请注意防暑降温。\n"
        "怎么做: 1)午后(11-16点)减少户外活动；"
        "2)户外工作者每小时到阴凉处休息10分钟；"
        "3)敏感人群(老人、儿童、慢病患者)减少长时间户外停留；"
        "4)多饮水，避免含咖啡因或酒精的饮品。\n"
        "请关注官方气象预报。"
    ),
}
WARNING_TEMPLATES_COLD = {
    "极高风险": (
        "【低温健康红色预警-紧急】{city}在{period}健康风险极高(评分{score}/5)。"
        "最低气温达{tmin}C，风寒指数{wc}C，{cold_std}。\n"
        "为什么危险: 极端低温可导致失温症、冻伤，心脑血管疾病急性发作风险骤增。"
        "Gasparrini 2015 Lancet: 低温归因死亡负担(7.29%)远高于高温(0.42%)。\n"
        "怎么做: 1)停止非必要户外活动；2)开放取暖场所，对独居老人、"
        "流浪人员实施安置救助；3)注意防寒保暖，防范水管冻裂与一氧化碳中毒；"
        "4)出现寒战、意识模糊:立即回暖并就医。\n"
        "预警生效至本周期结束。"
    ),
    "高风险": (
        "【低温健康橙色预警】{city}在{period}存在较高低温健康风险(评分{score}/5)。"
        "预计最低气温{tmin}C，风寒指数{wc}C。\n"
        "为什么危险: 严寒天气心脑血管疾病急性发作风险升高，"
        "呼吸道感染风险增加(低温低湿环境利于病毒传播, Resp Med 2009)。\n"
        "怎么做: 1)注意防寒保暖，尤其头部和四肢；2)取暖时注意通风，"
        "防止一氧化碳中毒；3)老年人减少外出，慢病患者按时服药；"
        "4)户外人员防冻伤，减少暴露时间。\n"
        "请关注官方气象预报。"
    ),
    "中风险": (
        "【低温健康黄色提示】{city}在{period}存在中度低温健康风险(评分{score}/5)。"
        "预计最低气温{tmin}C。\n"
        "怎么做: 1)注意保暖；2)心脑血管及呼吸系统疾病人群减少外出；"
        "3)室内取暖注意通风。"
    ),
}
WARNING_TEMPLATES_LOW = {
    "低风险": (
        "【气候健康提示】{city}在{period}气候健康风险较低(评分{score}/5)。"
        "天气条件总体安全，建议维持日常健康防护，关注官方气象预报。"
    ),
    "极低风险": (
        "【常规】{city}在{period}气候健康风险低(评分{score}/5)。"
        "天气条件总体安全，建议维持日常健康生活方式。"
    ),
}


def _near_threshold_tip(metrics):
    """判断是否接近风险阈值(但未达热浪/寒潮标准), 返回提示或空串。"""
    tmax = metrics.get("temperature_max_c")
    hi = metrics.get("heat_index_max_c")
    tmin = metrics.get("temperature_min_c")
    tips = []
    if tmax is not None and 33.0 <= tmax < 35.0:
        tips.append(f"最高温{tmax:.1f}C已接近35C高温阈值，午后减少户外活动")
    if hi is not None and 32.2 <= hi < 41.1:
        tips.append(f"热指数{hi:.1f}C达NOAA'警戒'级(>=32.2C)，注意体感闷热")
    if tmin is not None and -5.0 <= tmin < 0.0:
        tips.append(f"最低温{tmin:.1f}C接近0C，早晚注意保暖")
    return "。".join(tips)


def build_warning_text(metrics, city, period):
    """基于风险等级与风险源生成预警文案(只引用指标中的数值)。
    模板结构: 风险结论->为什么危险->具体行动->何时解除, 参考WHO热浪预警指南。
    """
    level = metrics["risk_level"]
    score = metrics["risk_score"]
    tmax = metrics.get("temperature_max_c") or "-"
    tmin = metrics.get("temperature_min_c") or "-"
    wc = metrics.get("wind_chill_min_c") or "-"
    hi = metrics.get("heat_index_max_c") or "-"
    hw_days = metrics.get("heatwave", {}).get("cma_days", 0)
    cold_std = metrics.get("coldwave", {}).get("standard_used", "寒潮")

    is_cold = metrics.get("coldwave", {}).get("active")
    is_heat = metrics.get("heatwave", {}).get("active")

    if is_cold and level in WARNING_TEMPLATES_COLD:
        text = WARNING_TEMPLATES_COLD[level].format(
            city=city, period=period, score=score, tmin=tmin, wc=wc, cold_std=cold_std)
    elif is_heat and level in WARNING_TEMPLATES_HEAT:
        text = WARNING_TEMPLATES_HEAT[level].format(
            city=city, period=period, score=score, tmax=tmax, hi=hi, hw_days=hw_days)
    else:
        text = WARNING_TEMPLATES_LOW.get(
            level, WARNING_TEMPLATES_LOW["低风险"]).format(
            city=city, period=period, score=score)
        # 临界提示: 接近高温/低温阈值但未触发事件
        tip = _near_threshold_tip(metrics)
        if tip:
            text += f" 注意: {tip}。"

    # 附加: 空气污染叠加提示
    aq = metrics.get("air_quality", {})
    if aq.get("aqi_label") in ("中度污染", "重度污染", "严重污染"):
        text += "\n注意: 空气质量较差，敏感人群请减少户外暴露并佩戴口罩。"
    elif aq.get("aqi_label") == "轻度污染":
        text += "\n提示: 空气轻度污染，敏感人群减少长时间户外停留。"
    return text


def build_report(metrics, city, period, latitude, longitude):
    """组装最终报告 JSON。"""
    # 重点人群 (按风险源定制: 高温只写热相关, 低温只写冷相关)
    # 无人口数据, 不写人口规模/分布相关表述, 仅列脆弱人群与通用健康建议
    is_cold = metrics.get("coldwave", {}).get("active")
    is_heat = metrics.get("heatwave", {}).get("active")
    hi_now = metrics.get("heat_index_max_c")
    tmax_now = metrics.get("temperature_max_c")

    if is_heat or (hi_now is not None and hi_now >= 32.2) or (tmax_now is not None and tmax_now >= 33):
        # 高温场景
        at_risk = [
            {"group": "老年人(65+)",
             "reason": "体温调节能力下降，高温相关死亡脆弱性显著(Sci Adv 2025; WHO)",
             "action": "高温天减少外出，保持室内降温，注意补水(WHO建议)"},
            {"group": "儿童(<5岁)",
             "reason": "体温调节系统未发育完全，高温下脱水风险高(WHO)",
             "action": "避免正午户外活动，注意补水(WHO建议)"},
            {"group": "户外工作者",
             "reason": "高温暴露时间长，热射病风险高(WHO/OSHA)",
             "action": "错峰作业，定时休息补水，配备防暑物资(WHO建议)"},
            {"group": "心脑血管/呼吸系统慢病患者",
             "reason": "高温诱发疾病急性发作(Gasparrini 2015 Lancet)",
             "action": "按时服药，症状加重及时就医"},
            {"group": "孕妇",
             "reason": "孕期体温调节负担增大(WHO)",
             "action": "避免高温环境，减少外出"},
        ]
    elif is_cold:
        # 低温场景
        at_risk = [
            {"group": "老年人(65+)",
             "reason": "低温死亡负担显著高于高温(Gasparrini 2015 Lancet)，保暖能力下降",
             "action": "低温天注意保暖，保持室内温度，减少外出"},
            {"group": "婴幼儿",
             "reason": "体温调节未发育完全，低温失温风险高",
             "action": "注意保暖，避免长时间户外暴露"},
            {"group": "户外工作者",
             "reason": "严寒暴露时间长，冻伤风险高(OSHA)",
             "action": "减少暴露时间，注意防冻保暖"},
            {"group": "心脑血管/呼吸系统慢病患者",
             "reason": "低温诱发心脑血管疾病急性发作(Resp Med 2009)",
             "action": "按时服药，注意呼吸道防护，症状加重及时就医"},
        ]
    else:
        # 温和天气
        at_risk = [
            {"group": "老年人(65+)",
             "reason": "温度敏感人群(WHO)",
             "action": "保持常规健康防护"},
            {"group": "儿童(<5岁)",
             "reason": "体温调节未发育完全(WHO)",
             "action": "保持常规健康防护"},
            {"group": "心脑血管/呼吸系统慢病患者",
             "reason": "温度变化敏感(Gasparrini 2015)",
             "action": "按时服药，关注天气变化"},
        ]

    # 不确定性说明 (参考 methodology.md §6 + 文献量化)
    uncertainty_parts = []
    if metrics.get("temperature_max_c") is not None:
        uncertainty_parts.append(
            "温度来自ERA5再分析(Hersbach 2020, doi:10.1002/qj.3803)，与地面观测偏差<1C；"
            "空间分辨率0.25，小于网格的社区被平滑")
    if metrics.get("heat_index_max_c") is not None:
        uncertainty_parts.append(
            "热指数为NOAA Rothfusz 1990经验公式近似值，高温高湿下误差增大")
    if metrics.get("wbgt_max_c") is not None:
        uncertainty_parts.append(
            "WBGT为BoM简化公式(无黑球温度)，户外直射下偏低1-2C")
    aq = metrics.get("air_quality", {})
    if aq.get("aqi_value") is not None:
        uncertainty_parts.append(
            "空气质量为CAMS再分析数据(融合卫星+地面站)；"
            "PM2.5每升10ug/m3总死亡风险+0.68%(Liu 2019 NEJM, 1480引)")
    uncertainty_parts.append(
        "J型曲线为示意性近似(Gasparrini 2015趋势)，真实曲线基于DLNM模型，"
        "此处不用于精确定量；风险评分权重(热冷40%+空气30%+脆弱30%)为复合矩阵设计")
    uncertainty_parts.append(
        "脆弱性数据为省级替代(第七次人口普查2020)，缺市级/社区级数据；"
        "本系统未纳入空调可及性、绿地覆盖率等适应资源数据")
    if not uncertainty_parts:
        uncertainty_parts.append("数据覆盖不完整，本报告结论仅具参考价值")

    # 公平性检查 (只声明实际有的和实际缺的, 不编造建议)
    vuln_score = metrics.get("vulnerability_score", 50)
    fairness = {
        "status": "已覆盖重点人群" if at_risk else "无重点人群提示",
        "covered_groups": [g["group"] for g in at_risk],
        "vulnerable_assessment": (
            f"脆弱性评分 {vuln_score}/100，基于第七次全国人口普查省级65+老年人口占比计算。"
            "本系统未纳入以下数据: 低收入社区分布、流动人口数据、空调可及性、绿地覆盖率。"
            "上述数据缺失可能导致弱势社区(老旧小区、流动人口聚集区)的风险被低估。"
        ),
        "data_gaps": [
            "低收入社区分布数据: 未纳入",
            "流动人口数据: 未纳入",
            "空调可及性(ac_proxy): 未纳入",
            "绿地覆盖率(green_ratio): 未纳入",
            "市级/社区级老年人口数据: 仅省级替代",
            "纳凉点/取暖点分布: 未纳入",
        ],
        "red_flags": [],
    }
    # 红旗检查 (只基于实际数据)
    if vuln_score > 60 and metrics.get("risk_score", 0) < 2:
        fairness["red_flags"].append(
            "高脆弱性区域风险评分偏低，可能因数据缺失导致低估")
    if not aq.get("aqi_value"):
        fairness["red_flags"].append(
            "空气质量数据缺失，用先验值替代，可能低估污染暴露")

    return {
        "report_version": "1.0",
        "generated_at": datetime.now().astimezone().isoformat(),
        "location": {"city": city, "latitude": latitude, "longitude": longitude},
        "period": {"start": period.split("~")[0], "end": period.split("~")[1]} if "~" in period else {"period_text": period},
        "metrics": metrics,
        "risk_summary": {
            "risk_score": metrics["risk_score"],
            "risk_level": metrics["risk_level"],
            "primary_hazard": metrics["primary_hazard"],
        },
        "at_risk_groups": [g["group"] for g in at_risk],
        "at_risk_details": at_risk,
        "warning_text": build_warning_text(metrics, city, period),
        "recommended_actions": metrics["recommended_actions"],
        "uncertainty": {
            "sources": uncertainty_parts,
        "statement": "本报告所有数值均来自公开数据源(Open-Meteo ERA5/CAMS)与科学公式，"
                     "详见证据链章节。数据缺口和公式近似已如实声明，未编造任何数据。",
        },
        "fairness_check": fairness,
        "evidence": metrics["evidence"],
        "references": [
            {"id": "lancet2015", "cite": "Gasparrini et al. 2015, The Lancet, doi:10.1016/S0140-6736(14)62114-0"},
            {"id": "nejm2019", "cite": "Liu et al. 2019, NEJM, doi:10.1056/NEJMoa1817364"},
            {"id": "pnas2010", "cite": "Sherwood & Huber 2010, PNAS, doi:10.1073/pnas.0913352107"},
            {"id": "whoaqg2021", "cite": "WHO 2021 AQG, https://www.who.int/publications/i/item/9789240034228"},
            {"id": "gbd2019", "cite": "GBD 2019 Risk Factors, The Lancet, doi:10.1016/S0140-6736(20)30752-2"},
            {"id": "envint2023", "cite": "Environ Int 2023, 482 cities, doi:10.1016/j.envint.2023.107825"},
        ],
    }


def build_word(report, out_path):
    """生成 Word 报告 (论文格式: 宋体正文/黑体标题/嵌入风险图)。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    m = report["metrics"]

    # 页面边距 (论文 A4 常规)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    # 正文: 宋体 12pt (小四), 西文 Times New Roman
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    def _set_cn(run, cn_font="宋体", size=12, bold=False, color=None):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _heading(text, level=1):
        # 黑体标题 (论文规范)
        h = doc.add_paragraph()
        if level == 1:
            _set_cn(h.add_run(text), "黑体", 15, True, (0x1A, 0x3A, 0x5C))
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        else:
            _set_cn(h.add_run(text), "黑体", 12, True, (0x1A, 0x3A, 0x5C))
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
        return h

    def _para(text, size=12, bold=False, cn="宋体"):
        p = doc.add_paragraph()
        _set_cn(p.add_run(text), cn, size, bold)
        return p

    # 主标题: 黑体 16pt 居中
    ht = doc.add_paragraph()
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn(ht.add_run(f"{report['location']['city']} 气候健康风险预警报告"), "黑体", 16, True)
    ht.paragraph_format.space_after = Pt(4)

    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn(sub.add_run(
        f"评估期: {report.get('period',{}).get('period_text','')}    "
        f"风险等级: {report['risk_summary']['risk_level']}"
        f"({report['risk_summary']['risk_score']}/5)    "
        f"主要风险: {report['risk_summary']['primary_hazard']}"), "楷体", 11)
    sub.paragraph_format.space_after = Pt(10)

    # 一、风险图 (嵌入)
    _heading("一、风险空间分布")
    map_png = os.path.join(os.path.dirname(out_path), "risk_map.png")
    if not os.path.exists(map_png):
        map_png = os.path.join(os.path.dirname(out_path), "dashboard_thumbnail.png")
    if os.path.exists(map_png):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(map_png, width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cn(cap.add_run(
            f"图1  {report['location']['city']}风险空间分布"
            "(城市网格, 绿=低风险 橙=中风险 红=高风险)"), "宋体", 9)
    else:
        _para("(风险图未生成)")

    # 二、预警文案
    _heading("二、预警文案")
    _para(report.get("warning_text", ""))

    # 三、核心指标
    _heading("三、核心气象与健康指标")
    aq = m.get("air_quality", {})
    table = doc.add_table(rows=8, cols=2, style="Table Grid")
    cells = [
        ("最高气温", f"{m.get('temperature_max_c','-')} C"),
        ("最低气温", f"{m.get('temperature_min_c','-')} C"),
        ("平均湿度", f"{m.get('humidity_mean_pct','-')} %"),
        ("热指数(最高)", f"{m.get('heat_index_max_c','-')} C"),
        ("WBGT(最高)", f"{m.get('wbgt_max_c','-')} C"),
        ("AQI", f"{aq.get('aqi_value','-')} ({aq.get('aqi_label','-')}, 首要:{aq.get('aqi_primary_pollutant','-')})"),
        ("热浪", f"{'是' if m.get('heatwave',{}).get('active') else '否'} ({m.get('heatwave',{}).get('standard_used','')})"),
        ("寒潮/严寒", f"{'是' if m.get('coldwave',{}).get('active') else '否'} ({m.get('coldwave',{}).get('standard_used','')})"),
    ]
    for i, (k, v) in enumerate(cells):
        c0 = table.cell(i, 0).paragraphs[0]
        c1 = table.cell(i, 1).paragraphs[0]
        _set_cn(c0.add_run(k), "黑体", 11)
        _set_cn(c1.add_run(str(v)), "宋体", 11)

    # 四、健康影响解读 (用文献 RR 数字解读当前数据)
    _heading("四、健康影响解读(文献量化)")
    m2 = report["metrics"]
    hi2 = m2.get("heat_index_max_c")
    pm25 = m2.get("air_quality", {}).get("pollutants", {}).get("pm2_5", {}).get("mean")
    tmax2 = m2.get("temperature_max_c")
    tmean2 = m2.get("temperature_mean_c")
    lines = []
    if tmax2 is not None and tmax2 >= 35:
        lines.append(f"最高温{tmax2:.1f}C达高温标准(>=35C)。"
                     "中国气象局标准: 连续3天>=35C为高温热浪，健康风险显著升高")
    if hi2 is not None and hi2 >= 41.1:
        lines.append(f"热指数{hi2:.1f}C达NOAA'危险'级(41.1C)，中暑风险显著升高，需警惕热射病")
    elif hi2 is not None and hi2 >= 32.2:
        lines.append(f"热指数{hi2:.1f}C达NOAA'警戒'级(32.2C)，体感闷热，敏感人群需减少午后户外")
    if pm25 is not None:
        excess = (pm25 - 10) / 10 * 0.68 if pm25 > 10 else 0
        lines.append(f"PM2.5均值{pm25:.1f}ug/m3。"
                     "依据Liu 2019 NEJM(652城市): PM2.5每升10ug/m3，总死亡风险+0.68%。"
                     f"相对参考浓度10ug/m3，估算超额死亡风险约{excess:.1f}%")
    if tmean2 is not None:
        mmt = 24.0
        if tmean2 > mmt:
            coef = 0.025
        else:
            coef = 0.010
        rr = math.exp(coef * (tmean2 - mmt) ** 2)
        lines.append(f"日均温{tmean2:.1f}C。基于Gasparrini 2015 J型曲线(MMT=24C)估算"
                     f"相对死亡风险RR约{rr:.2f}，{'低于' if rr<1.05 else '略高于'}最适温度")
    if not lines:
        lines.append("当前气象条件温和，无显著健康风险")
    for ln in lines:
        _para(f"- {ln}", 12)

    # 五、重点人群
    _heading("五、重点人群提示")
    _para("以下人群对气候健康风险更脆弱(依据: WHO气候与健康事实页, "
          "Gasparrini 2015 Lancet, Wang 2025 Sci Adv)。"
          "行动建议为WHO通用指南，非本系统特有。")
    for g in report.get("at_risk_details", []):
        _heading(g["group"], 2)
        _para(f"风险原因: {g['reason']}")
        _para(f"建议行动: {g['action']}")

    # 六、行动建议
    _heading("六、行动建议")
    for a in report.get("recommended_actions", []):
        _para(f"- {a}", 12)

    # 七、不确定性说明
    _heading("七、不确定性与置信度")
    _para(report.get("uncertainty", {}).get("statement", ""))
    for s in report.get("uncertainty", {}).get("sources", []):
        _para(f"- {s}", 12)

    # 八、公平性检查
    _heading("八、公平性检查")
    f = report.get("fairness_check", {})
    _para(f"总体状态: {f.get('status','')}")
    _heading("脆弱性评估", 2)
    _para(f.get("vulnerable_assessment", ""))
    _heading("数据缺口声明", 2)
    for gap in f.get("data_gaps", []):
        _para(f"- {gap}", 12)
    if f.get("red_flags"):
        _heading("红旗清单", 2)
        for rf in f["red_flags"]:
            _para(f"- {rf}", 12)

    # 九、证据链
    _heading("九、证据链")
    _para("本报告所有数值的计算方法与文献依据:")
    for e in report.get("evidence", []):
        _para(f"- {e['metric']}: {e['method']} - {e['source']}", 11)

    doc.save(out_path)


def build_markdown(report):
    """生成 Markdown 版报告（演示/传播用）。"""
    m = report["metrics"]
    lines = [
        f"# {report['location']['city']} 气候健康风险预警报告",
        "",
        f"> 生成时间: {report['generated_at']}",
        f"> 综合风险等级: **{report['risk_summary']['risk_level']}** "
        f"({report['risk_summary']['risk_score']}/5)",
        f"> 主要风险源: {report['risk_summary']['primary_hazard']}",
        "",
        "## 一、核心指标",
        "",
        "| 指标 | 数值 |",
        "|---|---|",
        f"| 最高气温 | {m.get('temperature_max_c')} °C |",
        f"| 最低气温 | {m.get('temperature_min_c')} °C |",
        f"| 平均湿度 | {m.get('humidity_mean_pct')} % |",
        f"| 热指数(最高) | {m.get('heat_index_max_c')} °C |",
        f"| WBGT(最高) | {m.get('wbgt_max_c')} °C |",
        f"| 风寒指数(最低) | {m.get('wind_chill_min_c')} °C |",
        f"| 热浪 | {'是' if m['heatwave']['active'] else '否'} ({m['heatwave']['standard_used']}) |",
        f"| 寒潮/严寒 | {'是' if m['coldwave']['active'] else '否'} ({m['coldwave']['standard_used']}) |",
        f"| AQI | {m['air_quality'].get('aqi_value')} ({m['air_quality'].get('aqi_label')}, "
        f"首要: {m['air_quality'].get('aqi_primary_pollutant')}) |",
        f"| PM2.5 | {m['air_quality']['pollutants'].get('pm2_5', {}).get('mean')} μg/m³ "
        f"({m['air_quality']['pollutants'].get('pm2_5', {}).get('who_level', '数据缺失')}) |",
        "",
        "## 二、预警文案",
        "",
        report["warning_text"],
        "",
        "## 三、重点人群与行动建议",
        "",
        "- 重点人群: " + ", ".join(report["at_risk_groups"]),
        "",
    ]
    lines += [f"- {a}" for a in report["recommended_actions"]]
    lines += ["", "## 四、不确定性", ""]
    lines += [f"- {u}" for u in report["uncertainty"]["sources"]]
    lines += ["", "## 五、证据链", ""]
    for e in report["evidence"]:
        lines.append(f"- `{e['metric']}`: {e['method']} — {e['source']}")
    return "\n".join(lines)


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    ap = argparse.ArgumentParser(description="生成气候健康风险预警报告")
    ap.add_argument("indices_json", help="compute_indices.py 输出")
    ap.add_argument("--city", default="未知城市")
    ap.add_argument("--lat", type=float, default=None)
    ap.add_argument("--lon", type=float, default=None)
    ap.add_argument("--period", default="评估期")
    ap.add_argument("--out-dir", default="./output")
    ap.add_argument("--map", action="store_true", help="生成 folium 风险地图")
    args = ap.parse_args()

    with open(args.indices_json, encoding="utf-8") as f:
        metrics = json.load(f)

    os.makedirs(args.out_dir, exist_ok=True)
    report = build_report(metrics, args.city, args.period, args.lat, args.lon)

    # JSON 报告
    report_path = os.path.join(args.out_dir, "report.json")
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"OK: {report_path}")

    # Markdown 报告
    md_path = os.path.join(args.out_dir, "report.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(build_markdown(report))
    print(f"OK: {md_path}")

    # 大屏 HTML + 静态风险图 (先于 Word, 让 Word 能嵌入 risk_map.png)
    if args.map:
        build_map(report, metrics, args, args.lat or 31.0, args.lon or 121.0,
                  args.out_dir)

    # Word 报告
    try:
        docx_path = os.path.join(args.out_dir, "report.docx")
        build_word(report, docx_path)
        print(f"OK: {docx_path}")
    except Exception as exc:
        print(f"警告: Word报告生成失败: {exc}")


def build_map(report, metrics, args, lat, lon, out_dir):
    """生成大屏 HTML (沿行政区边界裁剪色斑图 + KPI + 图表)。

    输出: dashboard.html (单文件, 含 4 张 tab 色斑图)
    数据: grid_risk 网格风险 + ERA5/CAMS 原始数据
    """
    import os
    import sys

    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from grid_risk import compute_grid

    print("  计算城市网格风险点 ...")
    vuln = report["metrics"].get("vulnerability_score", 50.0)
    aqi_label = report["metrics"].get("air_quality", {}).get("aqi_label") or "良"
    start = report.get("period", {}).get("start")
    end = report.get("period", {}).get("end")
    if not start or not end:
        print("警告: 缺少日期范围，跳过大屏")
        return
    points, ok = compute_grid(lat, lon, start, end, n=4,
                              vulnerability=vuln, aqi_label=aqi_label)
    print(f"  网格点: {ok}/{len(points)} 成功")

    city = report["location"]["city"]
    _skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    geojson_path = os.path.join(_skill_dir, "data", "geojson", f"{city}.json")
    if not os.path.exists(geojson_path):
        print(f"  警告: 无边界文件 {geojson_path}，跳过大屏")
        return

    # 保存网格点数据
    points_path = os.path.join(out_dir, "risk_grid_points.json")
    with open(points_path, "w", encoding="utf-8") as f:
        json.dump(points, f, ensure_ascii=False)

    # 加载原始 ERA5/AQI 数据
    era5_path = os.path.join(_skill_dir, "output", "tmp",
                             f"era5_{city}_{start}_{end}.json")
    aqi_path = os.path.join(_skill_dir, "output", "tmp",
                            f"aqi_{city}_{start}_{end}.json")
    era5_data = json.load(open(era5_path, encoding="utf-8")) if os.path.exists(era5_path) else {}
    aqi_data = json.load(open(aqi_path, encoding="utf-8")) if os.path.exists(aqi_path) else {}

    # 生成大屏 HTML
    try:
        from make_dashboard import build
        # 获取当天实时数据 + 近 30 天基线 + 异常判断 + 风险趋势
        realtime = {}
        try:
            sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
            from fetch_realtime import (fetch_today, fetch_baseline, anomaly_assessment,
                                        fetch_history_and_forecast, daily_risk_score)
            rt_today = fetch_today(lat, lon)
            rt_baseline = fetch_baseline(lat, lon, days=30)
            # 异常判断用当天日最高温 (非当前时刻温度), 符合气象预警标准
            hourly_temps = rt_today.get("hourly", {}).get("temperature_2m", [])
            today_tmax = max([t for t in hourly_temps if t is not None], default=None)
            rt_anomaly = anomaly_assessment(today_tmax, rt_baseline)
            # 风险趋势: 近15天历史 + 未来7天预报
            rt_trend = fetch_history_and_forecast(lat, lon, hist_days=15, forecast_days=7)
            rt_trend["risk_scores"] = [
                daily_risk_score(rt_trend["tmean"][i], rt_trend["tmax"][i],
                                 rt_trend["tmin"][i])
                for i in range(len(rt_trend["dates"]))
            ]
            realtime = {"today": rt_today, "baseline": rt_baseline,
                        "anomaly": rt_anomaly, "trend": rt_trend}
            print(f"  实时: 今日Tmax={today_tmax}C | 异常: {rt_anomaly['status']} | "
                  f"趋势: {len(rt_trend['dates'])}天")
        except Exception as exc:
            print(f"  警告: 实时数据获取失败({exc})，使用区间数据")

        out_html = os.path.join(out_dir, "dashboard.html")
        build(points, era5_data, aqi_data, report, report["metrics"],
              geojson_path, city, f"{start}~{end}", out_html, realtime=realtime)

        # 生成静态风险图 (Word 嵌入用)
        try:
            from make_risk_image import draw_risk_image
            draw_risk_image(geojson_path, points, city, start, end,
                            os.path.join(out_dir, "risk_map.png"))
        except Exception as exc:
            print(f"  警告: 静态风险图生成失败: {exc}")
    except Exception as exc:
        print(f"警告: 大屏生成失败: {exc}")


if __name__ == "__main__":
    main()
